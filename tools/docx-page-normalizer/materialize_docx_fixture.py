from __future__ import annotations

from datetime import datetime, timezone
from io import BytesIO
from pathlib import Path
import tempfile
import zipfile

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Inches, Pt


TOOL_ROOT = Path(__file__).resolve().parent
REPO_ROOT = TOOL_ROOT.parents[1]
SOURCE_PNG = REPO_ROOT / "eval" / "visual-page-normalization" / "cases" / "junior-readable-measurement.normalized.png"
OUTPUT_DOCX = REPO_ROOT / "eval" / "docx-page-normalization" / "cases" / "junior-image-backed-page.docx"
FIXED_TIME = datetime(2026, 7, 31, 14, 30, tzinfo=timezone.utc)
ZIP_TIME = (2026, 7, 31, 14, 30, 0)


def build_docx(output_path: Path) -> None:
    document = Document()
    core = document.core_properties
    core.title = "Public synthetic image-backed DOCX normalization fixture"
    core.subject = "compact_reference_guide; image-only technical fixture override"
    core.author = "Classroom Answer Toolkit"
    core.last_modified_by = "Classroom Answer Toolkit"
    core.created = FIXED_TIME
    core.modified = FIXED_TIME
    core.revision = 1

    section = document.sections[0]
    section.start_type = WD_SECTION.NEW_PAGE
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.right_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    normal = document.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    normal.element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal.element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")

    paragraph = document.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.paragraph_format.line_spacing = 1
    run = paragraph.add_run()
    run.add_picture(BytesIO(SOURCE_PNG.read_bytes()), width=Inches(6.5))

    output_path.parent.mkdir(parents=True, exist_ok=True)
    with tempfile.TemporaryDirectory(prefix="docx-image-backed-") as directory:
        raw_path = Path(directory) / "raw.docx"
        document.save(raw_path)
        with zipfile.ZipFile(raw_path) as source, zipfile.ZipFile(
            output_path,
            "w",
            compression=zipfile.ZIP_DEFLATED,
            compresslevel=9,
        ) as target:
            for info in sorted(source.infolist(), key=lambda item: item.filename):
                normalized = zipfile.ZipInfo(info.filename, ZIP_TIME)
                normalized.compress_type = zipfile.ZIP_DEFLATED
                normalized.create_system = 3
                normalized.external_attr = 0o100644 << 16
                target.writestr(normalized, source.read(info))


if __name__ == "__main__":
    build_docx(OUTPUT_DOCX)
    print(OUTPUT_DOCX)
